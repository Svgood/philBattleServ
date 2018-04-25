import docx
import mysql


if __name__ == '__main__':

    db = mysql.connector.connect(user="root", password="Svgood22",
                                 host="127.0.0.1", database="phibattle")

    cursor = db.cursor()

    doc = docx.Document("doc.docx")
    mas = []
    for i in range(1, 181):
        #print(doc.paragraphs[i].text)
        if i % 5 == 1:
            print(mas)
            mas = []
            q = doc.paragraphs[i].text[2:]
            if q[0] == '.':
                q = q[1:]
            mas.append(q)
        else:
            a = doc.paragraphs[i].text[2:]
            if a.find(";") != -1:
                a = a[:a.find(";")]

            if a.find(")") != -1:
                a = a[:a.find(")")] + a[a.find(")")+1:]

            if a[0] == ")":
                if len(mas) > 1:
                    mas.append(mas[1])
                    mas[1] = a[1:]
                else:
                    mas.append(a)
            else:
                mas.append(a)


    # for obj in mas:
    #     cursor.execute(
    #         """INSERT INTO questions (text, ans1, ans2, ans3, ans4, category)
    #         VALUES ('{}', '{}', '{}', '{}', '{}', 1);""".format(
    #             mas[0], mas[1], mas[2], mas[3], mas[4]))
    # db.commit()
    # cursor.close()
    # db.close()
